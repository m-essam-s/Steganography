from PIL import Image
from docx import Document
# Step 1: Convert the sentence to a binary stream
def sentence_to_binary(sentence):
    binary_stream = ""
    for char in sentence:
        # Check if the character is within the ASCII range
        if ord(char) < 128:
            # Convert each character to 7-bit ASCII
            ascii_code = ord(char) & 0x7F
            # Convert ASCII code to binary and append to the stream
            binary_stream += format(ascii_code, '07b')
        else:
            # Handle characters outside the ASCII range
            print("Character '{}' is outside the ASCII range and will be skipped.".format(char))
    return binary_stream

# Step 2: Choose an image and modify its pixels to hide the binary stream
def hide_binary_in_image(image_path, binary_stream):
    img = Image.open(image_path)
    width, height = img.size

    # Calculate the maximum number of pixels available for hiding the binary stream
    max_pixels = width * height

    # Check if the binary stream fits within the available pixels
    if len(binary_stream) > max_pixels * 3:  # Each pixel has 3 color channels (RGB)
        raise ValueError("Binary stream too long for the image size.")

    pixels = img.load()
    index = 0

    # Iterate through the image pixels and modify the least significant bit
    for y in range(height):
        for x in range(width):
            if index < len(binary_stream):
                # Get the pixel value
                pixel = list(img.getpixel((x, y)))
                # Modify the least significant bit of each color channel
                for i in range(3):  # RGB channels
                    if index < len(binary_stream):
                        pixel[i] &= ~1  # Clear the least significant bit
                        pixel[i] |= int(binary_stream[index])  # Set the bit according to the binary stream
                        index += 1
                # Update the pixel value
                img.putpixel((x, y), tuple(pixel))
            else:
                break

    # Save the modified image
    img.save("./_files/encoded_image.png")
    print("Image with hidden message saved as encoded_image.png")

# Step 4: Create a Word file containing the sentence, its integer, and binary ASCII code representations
def save_to_file(sentence, binary_stream):
    with open("./_files/sentence_details", "w") as file:
        file.write("Sentence: {}\n".format(sentence))
        file.write("Integer ASCII codes: {}\n".format(", ".join(str(ord(char)) for char in sentence)))
        file.write("Binary ASCII codes: {}\n".format(" ".join(format(ord(char), '07b') for char in sentence)))
        file.write("Binary Stream: {}".format(binary_stream))
        
# Create a Readable Word file containing the sentence, its integer, and binary ASCII code representations
def save_to_readable_word_file(sentence, binary_stream):
    from docx import Document  # Import Document class here

    document = Document()
    document.add_heading('Sentence Details', level=1)
    
    # Create a table to organize the sentence details
    table = document.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    # Add data to the table
    table.cell(0, 0).text = 'Sentence:'
    table.cell(0, 1).text = sentence
    table.cell(1, 0).text = 'Integer ASCII codes:'
    table.cell(1, 1).text = ', '.join(str(ord(char)) for char in sentence)
    table.cell(2, 0).text = 'Binary ASCII codes:'
    table.cell(2, 1).text = ' '.join(format(ord(char), '07b') for char in sentence)
    table.cell(3, 0).text = 'Binary Stream:'
    table.cell(3, 1).text = binary_stream
    
    # Save the document
    document.save("./_files/sentence_details.docx")
    print("Word file with sentence details saved as sentence_details.docx")
    
if __name__=="__main__":
    first_name="Mohamed"
    _id=221101060
    sentence = f"{first_name}_{_id}$"
    binary_stream = sentence_to_binary(sentence)
    print("Binary stream:", binary_stream)
    image_path = "./_files/me.jpg"  # Replace with the path to your image
    hide_binary_in_image(image_path, binary_stream)
    save_to_file(sentence, binary_stream)
    save_to_readable_word_file(sentence, binary_stream)
    print("Word file with sentence details saved as sentence_details.docx")
