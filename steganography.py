# implement a simple steganography technique to hide a sentence in an image and a text file
# The sentence is converted to a binary stream, which is then embedded in the least significant bits of the image pixels and text file spaces.
# The script also creates a Word document containing the sentence, its integer, and binary ASCII code representations.
# implemented by Mohamed Essam https://github.com/m-essam-s

from PIL import Image
from docx import Document
import os

# Step 1: Convert the sentence to a binary stream
def sentence_to_binary(sentence):
    binary_stream = ""
    for char in sentence:
        # Convert each character to 7-bit ASCII
        ascii_code = ord(char) & 0x7F
        # Convert ASCII code to binary and append to the stream
        binary_stream += format(ascii_code, '07b')
    return binary_stream

# Step 2: Choose an image and modify its pixels to hide the binary stream
def embed_bits_in_image(image_path, binary_stream):
    try:
        img = Image.open(image_path)
        if img.mode != 'RGB':
            img = img.convert('RGB')
        width, height = img.size

        # Check if the binary stream fits within the available pixels
        if len(binary_stream) > width * height * 3:  # Each pixel has 3 color channels (RGB)
            raise ValueError("Binary stream too long for the image size.")

        pixels = img.load()
        index = 0

        # Iterate through the image pixels and modify the least significant bit
        for y in range(height):
            for x in range(width):
                if index < len(binary_stream):
                    pixel = list(pixels[x, y])
                    for i in range(3):  # RGB channels
                        if index < len(binary_stream):
                            pixel[i] = (pixel[i] & ~1) | int(binary_stream[index])
                            index += 1
                    pixels[x, y] = tuple(pixel)
                else:
                    break

        # Save the modified image
        os.makedirs(os.path.dirname("./img_24_bit_depth_steganography/"), exist_ok=True)
        encoded_image_path = "./img_24_bit_depth_steganography/encoded_image.png"
        img.save(encoded_image_path)
        print(f"Image with hidden message saved as {encoded_image_path}")
    except Exception as e:
        print(f"Error embedding bits in image: {e}")

# Step 3: Embed the binary stream in a text file
def embed_bits_in_text(text_file_path, binary_stream):
    try:
        with open(text_file_path, "r") as file:
            text = file.read()

        result = []
        space_count = 0
        bit_length = len(binary_stream)
        
        i = 0
        while i < len(text):
            if text[i] == ' ':
                space_count += 1
                bit_index = (space_count - 1) % bit_length
                
                if binary_stream[bit_index] == '1':
                    result.append('  ')  # Replace space with double space for '1'
                else:
                    result.append(' ')  # Keep single space for '0'
            else:
                result.append(text[i])
            i += 1
        embedded_text = ''.join(result)
        
        os.makedirs(os.path.dirname("./text_file_steganography/"), exist_ok=True)
        with open("./text_file_steganography/text_after_steganography.txt", "w") as file:
            file.write(embedded_text)
        print("Text after steganography saved as ./text_file_steganography/text_after_steganography.txt")
    except Exception as e:
        print(f"Error embedding bits in text: {e}")

# Step 4: Create Text and Word files containing the sentence, its integer, and binary ASCII code representations
def save_to_file(sentence, binary_stream):
    try:
        os.makedirs("./sentence_details/", exist_ok=True)
        
        with open("./sentence_details/details.txt", "w") as file:
            file.write("Sentence: {}\n".format(sentence))
            file.write("Integer ASCII codes: {}\n".format(", ".join(str(ord(char)) for char in sentence)))
            file.write("Binary ASCII codes: {}\n".format(" ".join(format(ord(char), '07b') for char in sentence)))
            file.write("Binary Stream: {}".format(binary_stream))
        
        document = Document()
        document.add_heading('Sentence Details', level=1)
        
        # Create a table to organize the sentence details
        table = document.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        # Add data to the table
        table.cell(0, 0).text = 'Sentence:'
        table.cell(0, 1).text = sentence
        table.cell(1, 0).text = 'Integer ASCII codes:'
        table.cell(1, 1).text = ', '.join(str(ord(char)) for char in sentence)
        table.cell(2, 0).text = 'Binary ASCII codes:'
        table.cell(2, 1).text = ' '.join(format(ord(char), '07b') for char in sentence)
        table.cell(3, 0).text = 'Binary Stream:'
        table.cell(3, 1).text = binary_stream
        
        # Save the document
        document.save("./sentence_details/details.docx")
        print("Word file with sentence details saved as ./sentence_details/details.docx")
    except Exception as e:
        print(f"Error saving sentence details: {e}")

def main():
    first_name = "Mohamed"
    student_id = "221101060"
    sentence = f"{first_name}_{student_id}$"
    
    image_path = "./img_24_bit_depth_steganography/Cameraman_graycale.jpg"  # Replace with the path to your image
    text_file_path = "./text_file_steganography/Original_text.txt" # Replace with the path to your text file
    
    binary_stream = sentence_to_binary(sentence)
    
    embed_bits_in_image(image_path, binary_stream)
    embed_bits_in_text(text_file_path, binary_stream)
    save_to_file(sentence, binary_stream)
    
    print("--------------Sentence details--------------")
    with open("./sentence_details/details.txt", "r") as file:
        print(file.read())
        
    print("Text file with sentence details saved as ./sentence_details/details.txt")
    print("Word file with sentence details saved as ./sentence_details/details.docx")

if __name__ == "__main__":
    main()
